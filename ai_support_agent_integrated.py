import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import uuid
import time
import json
import sqlite3
from typing import List, Dict, Optional, Tuple
import re
from io import BytesIO
import base64

# For PDF/DOCX/Excel processing
try:
    import PyPDF2
    from docx import Document
    import openpyxl
    from openpyxl import load_workbook
except ImportError:
    st.warning("⚠️ Please install: pip install PyPDF2 python-docx openpyxl")

# For email functionality
try:
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    import imaplib
    import email
except ImportError:
    st.warning("⚠️ Email libraries not available")

# For WhatsApp integration
try:
    import webbrowser
    import urllib.parse
except ImportError:
    st.warning("⚠️ WhatsApp libraries not available")

# For vector database and embeddings
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_openai import OpenAIEmbeddings, ChatOpenAI
    from langchain.chains import RetrievalQA
    from langchain.docstore.document import Document as LangchainDocument
    from langchain.prompts import PromptTemplate
except ImportError:
    st.warning("⚠️ Please install: pip install langchain langchain-openai langchain-community faiss-cpu")

# For language detection and translation
try:
    from langdetect import detect, LangDetectException
    from deep_translator import GoogleTranslator
except ImportError:
    st.warning("⚠️ Please install: pip install langdetect deep-translator")

# For web scraping
try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    st.warning("⚠️ Please install: pip install requests beautifulsoup4")

# For OCR
try:
    import pytesseract
    from PIL import Image
except ImportError:
    st.warning("⚠️ Please install: pip install pytesseract Pillow")

# For semantic similarity (confidence scoring)
try:
    from sklearn.metrics.pairwise import cosine_similarity
    import numpy as np
except ImportError:
    st.warning("⚠️ Please install: pip install scikit-learn numpy")

# Page configuration
st.set_page_config(
    page_title="AI Support Agent Pro",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Enhanced Custom CSS with better color contrast
st.markdown("""
<style>
    /* Main theme colors with high contrast */
    :root {
        --primary-blue: #0066CC;
        --primary-dark: #003D7A;
        --success-green: #00A86B;
        --warning-orange: #FF8C00;
        --danger-red: #DC143C;
        --bg-light: #F8F9FA;
        --bg-white: #FFFFFF;
        --text-dark: #212529;
        --text-light: #6C757D;
        --border-color: #DEE2E6;
    }
    
    /* Header styling */
    .main-header {
        font-size: 2.8rem;
        font-weight: 800;
        background: linear-gradient(135deg, #0066CC 0%, #00A86B 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 2rem;
        padding: 1rem;
    }
    
    .sub-header {
        font-size: 1.5rem;
        font-weight: 600;
        color: var(--primary-dark);
        margin-bottom: 1rem;
        border-bottom: 3px solid var(--primary-blue);
        padding-bottom: 0.5rem;
    }
    
    /* Chat message styling with high contrast */
    .chat-message {
        padding: 1.2rem;
        border-radius: 12px;
        margin-bottom: 1rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        border: 1px solid var(--border-color);
    }
    
    .user-message {
        background: linear-gradient(135deg, #E3F2FD 0%, #BBDEFB 100%);
        margin-left: 3rem;
        border-left: 4px solid #0066CC;
        color: #003D7A;
    }
    
    .bot-message {
        background: linear-gradient(135deg, #F5F5F5 0%, #E8E8E8 100%);
        margin-right: 3rem;
        border-left: 4px solid #00A86B;
        color: #212529;
    }
    
    /* Confidence badges with high contrast */
    .confidence-badge {
        display: inline-block;
        padding: 4px 12px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
        margin: 0 4px;
    }
    
    .confidence-high {
        background-color: #00A86B;
        color: #FFFFFF;
    }
    
    .confidence-medium {
        background-color: #FF8C00;
        color: #FFFFFF;
    }
    
    .confidence-low {
        background-color: #DC143C;
        color: #FFFFFF;
    }
    
    /* Ticket card styling */
    .ticket-card {
        padding: 1.5rem;
        border: 2px solid var(--border-color);
        border-radius: 12px;
        margin-bottom: 1rem;
        background: var(--bg-white);
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
    }
    
    .ticket-card:hover {
        box-shadow: 0 6px 20px rgba(0,0,0,0.15);
        transform: translateY(-2px);
    }
    
    /* Priority badges */
    .priority-high {
        background-color: #DC143C;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    .priority-medium {
        background-color: #FF8C00;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    .priority-low {
        background-color: #00A86B;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 700;
        font-size: 0.85rem;
    }
    
    /* Status badges */
    .status-open {
        background-color: #0066CC;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    .status-progress {
        background-color: #FF8C00;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    .status-closed {
        background-color: #6C757D;
        color: white;
        padding: 6px 14px;
        border-radius: 20px;
        font-weight: 600;
    }
    
    /* Contact banner with excellent contrast */
    .contact-banner {
        background: linear-gradient(135deg, #0066CC 0%, #00A86B 100%);
        padding: 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 8px 24px rgba(0,0,0,0.15);
    }
    
    .contact-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        text-align: center;
        min-width: 280px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    
    .contact-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.2);
    }
    
    .contact-card h4 {
        color: var(--primary-dark);
        margin-top: 0;
        font-weight: 700;
    }
    
    .contact-card p {
        color: var(--text-dark);
        font-size: 1rem;
        font-weight: 500;
        margin: 10px 0;
    }
    
    /* Buttons with high contrast */
    .whatsapp-button {
        background-color: #25D366;
        color: white;
        padding: 12px 24px;
        border-radius: 8px;
        text-decoration: none;
        display: inline-block;
        margin: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        border: 2px solid #128C7E;
    }
    
    .whatsapp-button:hover {
        background-color: #128C7E;
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(37, 211, 102, 0.4);
    }
    
    .email-button {
        background-color: #0066CC;
        color: white;
        padding: 12px 24px;
        border-radius: 8px;
        text-decoration: none;
        display: inline-block;
        margin: 8px;
        font-weight: 600;
        transition: all 0.3s ease;
        border: 2px solid #003D7A;
    }
    
    .email-button:hover {
        background-color: #003D7A;
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(0, 102, 204, 0.4);
    }
    
    /* Metric cards with better visibility */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        border: 2px solid var(--border-color);
        text-align: center;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 800;
        color: var(--primary-blue);
    }
    
    .metric-label {
        font-size: 1rem;
        color: var(--text-light);
        font-weight: 600;
        margin-top: 0.5rem;
    }
    
    /* Info boxes with high contrast */
    .info-box {
        background: #E3F2FD;
        border-left: 5px solid #0066CC;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #003D7A;
        font-weight: 500;
    }
    
    .success-box {
        background: #D4EDDA;
        border-left: 5px solid #00A86B;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #155724;
        font-weight: 500;
    }
    
    .warning-box {
        background: #FFF3CD;
        border-left: 5px solid #FF8C00;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #856404;
        font-weight: 500;
    }
    
    .error-box {
        background: #F8D7DA;
        border-left: 5px solid #DC143C;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: #721C24;
        font-weight: 500;
    }
    
    /* Sidebar styling */
    .sidebar .sidebar-content {
        background-color: var(--bg-light);
    }
    
    /* Table styling */
    .dataframe {
        border: 2px solid var(--border-color) !important;
    }
    
    .dataframe th {
        background-color: var(--primary-blue) !important;
        color: white !important;
        font-weight: 700 !important;
    }
    
    .dataframe td {
        color: var(--text-dark) !important;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background-color: var(--bg-light);
        border: 2px solid var(--border-color);
        border-radius: 8px;
        font-weight: 600;
        color: var(--text-dark);
    }
    
    /* Footer */
    .footer {
        text-align: center;
        padding: 2rem;
        background: var(--bg-light);
        border-top: 3px solid var(--primary-blue);
        margin-top: 3rem;
        color: var(--text-dark);
    }
</style>
""", unsafe_allow_html=True)

# Database setup
def init_database():
    """Initialize SQLite database for persistence"""
    conn = sqlite3.connect('support_agent.db', check_same_thread=False)
    c = conn.cursor()
    
    # Chat history table
    c.execute('''CREATE TABLE IF NOT EXISTS chat_history
                 (id TEXT PRIMARY KEY, role TEXT, message TEXT, timestamp TEXT, 
                  confidence REAL, response_time REAL, language TEXT, category TEXT)''')
    
    # Tickets table
    c.execute('''CREATE TABLE IF NOT EXISTS tickets
                 (id TEXT PRIMARY KEY, query TEXT, language TEXT, category TEXT, 
                  status TEXT, priority TEXT, assigned_to TEXT, timestamp TEXT, 
                  resolved_at TEXT, resolution_time REAL, channel TEXT)''')
    
    # Feedback table
    c.execute('''CREATE TABLE IF NOT EXISTS feedback
                 (chat_id TEXT, feedback TEXT, timestamp TEXT, comment TEXT)''')
    
    # Analytics table
    c.execute('''CREATE TABLE IF NOT EXISTS analytics
                 (date TEXT, total_queries INTEGER, answered INTEGER, escalated INTEGER,
                  avg_response_time REAL, avg_confidence REAL)''')
    
    conn.commit()
    return conn

# Initialize database
if 'db_conn' not in st.session_state:
    st.session_state.db_conn = init_database()

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'tickets' not in st.session_state:
    st.session_state.tickets = []
if 'analytics' not in st.session_state:
    st.session_state.analytics = {
        'total_queries': 0,
        'answered': 0,
        'escalated': 0,
        'languages': {'English': 0, 'Hindi': 0, 'Marathi': 0, 'Other': 0},
        'categories': {'Billing': 0, 'Technical': 0, 'General': 0, 'Product': 0}
    }
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'knowledge_base_text' not in st.session_state:
    st.session_state.knowledge_base_text = ""
if 'feedback' not in st.session_state:
    st.session_state.feedback = {}
if 'embeddings_model' not in st.session_state:
    st.session_state.embeddings_model = None
if 'channel_stats' not in st.session_state:
    st.session_state.channel_stats = {
        'Website': 0,
        'WhatsApp': 0,
        'Email': 0,
        'Manual': 0
    }
if 'email_queue' not in st.session_state:
    st.session_state.email_queue = []
if 'whatsapp_messages' not in st.session_state:
    st.session_state.whatsapp_messages = []
if 'gmail_config' not in st.session_state:
    st.session_state.gmail_config = {
        'email': 'support@yourcompany.com',
        'smtp_server': 'smtp.gmail.com',
        'smtp_port': 587,
        'imap_server': 'imap.gmail.com'
    }
if 'whatsapp_config' not in st.session_state:
    st.session_state.whatsapp_config = {
        'phone_number': '+1234567890',
        'wa_link': 'https://wa.me/1234567890'
    }
if 'kb_processed' not in st.session_state:
    st.session_state.kb_processed = False

# Helper Functions
def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"❌ Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"❌ Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_excel(file) -> str:
    """Extract text from Excel file (XLSX/XLS)"""
    try:
        wb = load_workbook(file, read_only=True, data_only=True)
        text = ""
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"\n[Sheet: {sheet_name}]\n"
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) for cell in row if cell is not None]
                if row_data:
                    text += " | ".join(row_data) + "\n"
        
        wb.close()
        return text
    except Exception as e:
        st.error(f"❌ Error reading Excel file: {str(e)}")
        return ""

def extract_text_from_image(image_file) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"❌ Error performing OCR: {str(e)}")
        return ""

def scrape_url(url: str) -> str:
    """Scrape text content from URL"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        st.error(f"❌ Error scraping URL {url}: {str(e)}")
        return ""

def detect_language(text: str) -> str:
    """Detect language of text"""
    try:
        lang_code = detect(text)
        lang_map = {
            'en': 'English',
            'hi': 'Hindi',
            'mr': 'Marathi',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German'
        }
        return lang_map.get(lang_code, 'Other')
    except:
        return 'English'

def translate_text(text: str, target_lang: str) -> str:
    """Translate text to target language"""
    try:
        if target_lang == 'English':
            return text
        
        lang_code_map = {
            'Hindi': 'hi',
            'Marathi': 'mr',
            'Spanish': 'es',
            'French': 'fr',
            'German': 'de'
        }
        
        target_code = lang_code_map.get(target_lang, 'en')
        translator = GoogleTranslator(source='auto', target=target_code)
        translated = translator.translate(text)
        return translated
    except Exception as e:
        st.warning(f"⚠️ Translation failed: {str(e)}")
        return text

def categorize_query(query: str) -> str:
    """Enhanced keyword-based categorization"""
    query_lower = query.lower()
    
    billing_keywords = ['payment', 'invoice', 'bill', 'charge', 'refund', 'price', 'cost', 
                       'subscription', 'card', 'billing', 'money', 'paid', 'transaction']
    technical_keywords = ['error', 'bug', 'issue', 'problem', 'not working', 'broken', 
                         'crash', 'slow', 'loading', 'login', 'access', 'technical', 'password',
                         'installation', 'update', 'sync']
    product_keywords = ['feature', 'how to', 'tutorial', 'guide', 'demo', 'product', 
                       'functionality', 'use', 'work', 'setup', 'configure']
    
    billing_score = sum(1 for keyword in billing_keywords if keyword in query_lower)
    technical_score = sum(1 for keyword in technical_keywords if keyword in query_lower)
    product_score = sum(1 for keyword in product_keywords if keyword in query_lower)
    
    scores = {
        'Billing': billing_score,
        'Technical': technical_score,
        'Product': product_score
    }
    
    max_category = max(scores, key=scores.get)
    return max_category if scores[max_category] > 0 else 'General'

def assign_priority(confidence: float, category: str) -> str:
    """Assign priority based on confidence and category"""
    if confidence < 0.4 or category == 'Billing':
        return 'High'
    elif confidence < 0.6 or category == 'Technical':
        return 'Medium'
    else:
        return 'Low'

def create_vector_store(text: str, openai_api_key: str):
    """Create FAISS vector store from text with caching"""
    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        
        documents = [LangchainDocument(page_content=chunk) for chunk in chunks]
        
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        st.session_state.embeddings_model = embeddings
        vector_store = FAISS.from_documents(documents, embeddings)
        
        return vector_store
    except Exception as e:
        st.error(f"❌ Error creating vector store: {str(e)}")
        return None

def calculate_semantic_confidence(query: str, retrieved_docs: List, answer: str, embeddings) -> float:
    """Calculate confidence score using semantic similarity"""
    try:
        query_embedding = embeddings.embed_query(query)
        answer_embedding = embeddings.embed_query(answer)
        
        query_answer_sim = cosine_similarity(
            [query_embedding], 
            [answer_embedding]
        )[0][0]
        
        if retrieved_docs:
            doc_embeddings = [embeddings.embed_query(doc.page_content) for doc in retrieved_docs]
            doc_similarities = cosine_similarity([query_embedding], doc_embeddings)[0]
            avg_doc_sim = np.mean(doc_similarities)
        else:
            avg_doc_sim = 0.0
        
        confidence = (0.4 * query_answer_sim + 0.6 * avg_doc_sim)
        
        if len(answer.split()) < 10:
            confidence *= 0.7
        
        return float(min(max(confidence, 0.0), 1.0))
    except Exception as e:
        st.warning(f"⚠️ Confidence calculation error: {str(e)}")
        return 0.7 if len(retrieved_docs) > 0 and len(answer) > 50 else 0.4

def get_ai_response(query: str, vector_store, openai_api_key: str, target_language: str = 'English') -> Tuple[str, float, List]:
    """Get AI response using RAG with multilingual support"""
    try:
        english_query = query
        if target_language != 'English':
            try:
                translator = GoogleTranslator(source='auto', target='en')
                english_query = translator.translate(query)
            except:
                pass
        
        llm = ChatOpenAI(
            model_name="gpt-3.5-turbo",
            temperature=0.3,
            openai_api_key=openai_api_key
        )
        
        prompt_template = """You are a helpful and professional customer support assistant. 
        Use the following context to answer the question accurately and concisely. 
        If you cannot find the answer in the context, politely inform the customer and suggest contacting support for further assistance.
        
        Context: {context}
        
        Question: {question}
        
        Provide a helpful, accurate, and friendly answer:"""
        
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vector_store.as_retriever(search_kwargs={"k": 4}),
            return_source_documents=True,
            chain_type_kwargs={"prompt": PROMPT}
        )
        
        result = qa_chain({"query": english_query})
        
        answer = result['result']
        source_docs = result['source_documents']
        
        confidence = calculate_semantic_confidence(
            english_query, 
            source_docs, 
            answer, 
            st.session_state.embeddings_model
        )
        
        if target_language != 'English':
            answer = translate_text(answer, target_language)
        
        return answer, confidence, source_docs
    except Exception as e:
        st.error(f"❌ Error getting AI response: {str(e)}")
        error_msg = "I apologize, but I encountered an error processing your query. Please try again or contact support."
        if target_language != 'English':
            error_msg = translate_text(error_msg, target_language)
        return error_msg, 0.3, []

def create_ticket(query: str, language: str, category: str, confidence: float, channel: str = 'Website'):
    """Create escalation ticket with priority and assignment"""
    priority = assign_priority(confidence, category)
    
    agents = ['Agent A', 'Agent B', 'Agent C', 'Agent D']
    assigned_to = agents[len(st.session_state.tickets) % len(agents)]
    
    ticket = {
        'id': str(uuid.uuid4())[:8].upper(),
        'query': query,
        'language': language,
        'category': category,
        'status': 'Open',
        'priority': priority,
        'assigned_to': assigned_to,
        'channel': channel,
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'resolved_at': None,
        'resolution_time': None
    }
    
    st.session_state.tickets.append(ticket)
    st.session_state.analytics['escalated'] += 1
    
    conn = st.session_state.db_conn
    c = conn.cursor()
    c.execute('''INSERT INTO tickets VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
              (ticket['id'], ticket['query'], ticket['language'], ticket['category'],
               ticket['status'], ticket['priority'], ticket['assigned_to'], 
               ticket['timestamp'], ticket['resolved_at'], ticket['resolution_time'],
               ticket['channel']))
    conn.commit()
    
    return ticket['id']

def save_chat_to_db(chat_entry: Dict):
    """Save chat entry to database"""
    conn = st.session_state.db_conn
    c = conn.cursor()
    
    chat_id = str(uuid.uuid4())
    c.execute('''INSERT INTO chat_history VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
              (chat_id, chat_entry['role'], chat_entry['message'], 
               chat_entry['timestamp'].strftime("%Y-%m-%d %H:%M:%S"),
               chat_entry.get('confidence', None),
               chat_entry.get('response_time', None),
               chat_entry.get('language', None),
               chat_entry.get('category', None)))
    conn.commit()
    return chat_id

def update_daily_analytics():
    """Update daily analytics in database"""
    try:
        conn = st.session_state.db_conn
        c = conn.cursor()
        
        today = datetime.now().strftime("%Y-%m-%d")
        analytics = st.session_state.analytics
        
        bot_messages = [msg for msg in st.session_state.chat_history if msg.get('role') == 'bot']
        avg_response_time = sum(msg.get('response_time', 0) for msg in bot_messages) / len(bot_messages) if bot_messages else 0
        avg_confidence = sum(msg.get('confidence', 0) for msg in bot_messages) / len(bot_messages) if bot_messages else 0
        
        c.execute('''INSERT OR REPLACE INTO analytics VALUES (?, ?, ?, ?, ?, ?)''',
                  (today, analytics.get('total_queries', 0), analytics.get('answered', 0), 
                   analytics.get('escalated', 0), avg_response_time, avg_confidence))
        conn.commit()
    except Exception as e:
        st.error(f"❌ Error updating analytics: {str(e)}")

def send_gmail(to_email: str, subject: str, message: str, app_password: str) -> bool:
    """Send email via Gmail SMTP"""
    try:
        config = st.session_state.gmail_config
        
        msg = MIMEMultipart()
        msg['From'] = config['email']
        msg['To'] = to_email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(message, 'plain'))
        
        server = smtplib.SMTP(config['smtp_server'], config['smtp_port'])
        server.starttls()
        server.login(config['email'], app_password)
        server.send_message(msg)
        server.quit()
        
        return True
    except Exception as e:
        st.error(f"❌ Gmail sending failed: {str(e)}")
        return False

def check_gmail_inbox(app_password: str, max_emails: int = 10) -> List[Dict]:
    """Check Gmail inbox for new emails"""
    try:
        config = st.session_state.gmail_config
        
        mail = imaplib.IMAP4_SSL(config['imap_server'])
        mail.login(config['email'], app_password)
        mail.select('inbox')
        
        _, search_data = mail.search(None, 'UNSEEN')
        email_ids = search_data[0].split()[-max_emails:]
        
        emails = []
        for email_id in email_ids:
            _, data = mail.fetch(email_id, '(RFC822)')
            msg = email.message_from_bytes(data[0][1])
            
            subject = msg['subject']
            from_email = msg['from']
            
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True).decode()
                        break
            else:
                body = msg.get_payload(decode=True).decode()
            
            emails.append({
                'id': email_id.decode(),
                'from': from_email,
                'subject': subject,
                'body': body,
                'timestamp': datetime.now()
            })
        
        mail.close()
        mail.logout()
        
        return emails
    except Exception as e:
        st.error(f"❌ Error checking Gmail: {str(e)}")
        return []

def create_whatsapp_link(phone_number: str, message: str) -> str:
    """Create WhatsApp direct message link"""
    encoded_message = urllib.parse.quote(message)
    phone_clean = phone_number.replace('+', '').replace('-', '').replace(' ', '')
    return f"https://wa.me/{phone_clean}?text={encoded_message}"

def open_whatsapp_chat(message: str = ""):
    """Open WhatsApp chat with pre-filled message"""
    config = st.session_state.whatsapp_config
    wa_link = create_whatsapp_link(config['phone_number'], message)
    return wa_link

def process_multi_channel_query(query: str, channel: str, openai_api_key: str, 
                                contact_info: dict = None) -> tuple:
    """Process query from any channel and route appropriately"""
    
    if not st.session_state.vector_store:
        return "Knowledge base not loaded. Please upload documents first.", 0.3, None
    
    language = detect_language(query)
    category = categorize_query(query)
    
    st.session_state.channel_stats[channel] = st.session_state.channel_stats.get(channel, 0) + 1
    
    answer, confidence, source_docs = get_ai_response(
        query, 
        st.session_state.vector_store, 
        openai_api_key,
        language
    )
    
    ticket_id = None
    if confidence < 0.6:
        ticket_id = create_ticket(query, language, category, confidence, channel)
        escalation_msg = f"\n\n⚠️ This query has been escalated to our support team. Ticket ID: {ticket_id}"
        if language != 'English':
            escalation_msg = translate_text(escalation_msg, language)
        answer += escalation_msg
    
    return answer, confidence, ticket_id

def get_confidence_badge(confidence: float) -> str:
    """Generate HTML badge for confidence score"""
    if confidence >= 0.7:
        return f'<span class="confidence-badge confidence-high">🟢 High: {confidence:.1%}</span>'
    elif confidence >= 0.5:
        return f'<span class="confidence-badge confidence-medium">🟡 Medium: {confidence:.1%}</span>'
    else:
        return f'<span class="confidence-badge confidence-low">🔴 Low: {confidence:.1%}</span>'

def get_priority_badge(priority: str) -> str:
    """Generate HTML badge for priority"""
    return f'<span class="priority-{priority.lower()}">{priority}</span>'

def get_status_badge(status: str) -> str:
    """Generate HTML badge for status"""
    status_map = {
        'Open': 'open',
        'In Progress': 'progress',
        'Closed': 'closed'
    }
    return f'<span class="status-{status_map.get(status, "open")}">{status}</span>'

# Sidebar - Knowledge Management
with st.sidebar:
    st.markdown("### 📚 Knowledge Base")
    
    openai_api_key = st.text_input(
        "🔑 OpenAI API Key", 
        type="password", 
        help="Enter your OpenAI API key to enable AI features"
    )
    
    st.markdown("---")
    
    # Configuration section
    with st.expander("⚙️ System Configuration"):
        st.markdown("**📧 Email Settings**")
        gmail_email = st.text_input("Gmail Address", value=st.session_state.gmail_config['email'])
        
        st.markdown("**💬 WhatsApp Settings**")
        whatsapp_phone = st.text_input("WhatsApp Number", value=st.session_state.whatsapp_config['phone_number'])
        
        if st.button("💾 Save Configuration"):
            st.session_state.gmail_config['email'] = gmail_email
            st.session_state.whatsapp_config['phone_number'] = whatsapp_phone
            st.session_state.whatsapp_config['wa_link'] = f'https://wa.me/{whatsapp_phone.replace("+", "")}'
            st.success("✅ Configuration saved!")
    
    st.markdown("---")
    
    # File upload section
    st.markdown("**📁 Upload Knowledge Base**")
    uploaded_files = st.file_uploader(
        "Upload documents",
        type=['pdf', 'docx', 'xlsx', 'xls', 'png', 'jpg', 'jpeg'],
        accept_multiple_files=True,
        help="Supported: PDF, DOCX, Excel, Images (OCR)"
    )
    
    # URL input section
    url_input = st.text_area(
        "🌐 Or enter URLs (one per line)", 
        height=100,
        help="Enter website URLs to scrape content"
    )
    
    # Process button
    if st.button("🚀 Process Knowledge Base", type="primary", use_container_width=True):
        if not openai_api_key:
            st.error("❌ Please enter your OpenAI API key first!")
        else:
            with st.spinner("🔄 Processing knowledge base..."):
                all_text = ""
                
                # Process uploaded files
                if uploaded_files:
                    progress_bar = st.progress(0)
                    for idx, file in enumerate(uploaded_files):
                        file_type = file.name.split('.')[-1].lower()
                        st.info(f"📄 Processing: {file.name}")
                        
                        if file_type == 'pdf':
                            all_text += f"\n[PDF: {file.name}]\n{extract_text_from_pdf(file)}\n\n"
                        elif file_type == 'docx':
                            all_text += f"\n[DOCX: {file.name}]\n{extract_text_from_docx(file)}\n\n"
                        elif file_type in ['xlsx', 'xls']:
                            excel_text = extract_text_from_excel(file)
                            if excel_text:
                                all_text += f"\n[Excel: {file.name}]\n{excel_text}\n\n"
                        elif file_type in ['png', 'jpg', 'jpeg']:
                            ocr_text = extract_text_from_image(file)
                            if ocr_text:
                                all_text += f"\n[Image OCR: {file.name}]\n{ocr_text}\n\n"
                        
                        progress_bar.progress((idx + 1) / len(uploaded_files))
                    progress_bar.empty()
                
                # Process URLs
                if url_input.strip():
                    urls = [url.strip() for url in url_input.split('\n') if url.strip()]
                    progress_bar = st.progress(0)
                    for idx, url in enumerate(urls):
                        st.info(f"🌐 Scraping: {url[:50]}...")
                        scraped_text = scrape_url(url)
                        if scraped_text:
                            all_text += f"\n[URL: {url}]\n{scraped_text}\n\n"
                        progress_bar.progress((idx + 1) / len(urls))
                    progress_bar.empty()
                
                # Create vector store
                if all_text.strip():
                    st.session_state.knowledge_base_text = all_text
                    st.session_state.vector_store = create_vector_store(all_text, openai_api_key)
                    
                    if st.session_state.vector_store:
                        st.session_state.kb_processed = True
                        st.success(f"✅ Successfully processed {len(all_text):,} characters!")
                        st.balloons()
                    else:
                        st.error("❌ Failed to create vector store")
                else:
                    st.warning("⚠️ No content to process. Please upload files or enter URLs.")
    
    # Knowledge base status
    st.markdown("---")
    st.markdown("**📊 Knowledge Base Status**")
    
    if st.session_state.vector_store:
        st.success("✅ Active")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Size", f"{len(st.session_state.knowledge_base_text):,}")
        with col2:
            num_docs = st.session_state.vector_store.index.ntotal if st.session_state.vector_store else 0
            st.metric("Docs", num_docs)
    else:
        st.info("ℹ️ Not loaded")
    
    # Export section
    st.markdown("---")
    st.markdown("**📥 Export Data**")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💬 Chats", use_container_width=True):
            if st.session_state.chat_history:
                df = pd.DataFrame(st.session_state.chat_history)
                csv = df.to_csv(index=False)
                st.download_button(
                    "⬇️ Download",
                    csv,
                    "chat_history.csv",
                    "text/csv",
                    use_container_width=True
                )
    
    with col2:
        if st.button("🎫 Tickets", use_container_width=True):
            if st.session_state.tickets:
                df = pd.DataFrame(st.session_state.tickets)
                csv = df.to_csv(index=False)
                st.download_button(
                    "⬇️ Download",
                    csv,
                    "tickets.csv",
                    "text/csv",
                    use_container_width=True
                )
    
    # Clear data
    st.markdown("---")
    if st.button("🗑️ Clear All Data", type="secondary", use_container_width=True):
        if st.button("⚠️ Confirm Clear All", type="primary", use_container_width=True):
            st.session_state.chat_history = []
            st.session_state.tickets = []
            st.session_state.analytics = {
                'total_queries': 0,
                'answered': 0,
                'escalated': 0,
                'languages': {'English': 0, 'Hindi': 0, 'Marathi': 0, 'Other': 0},
                'categories': {'Billing': 0, 'Technical': 0, 'General': 0, 'Product': 0}
            }
            st.session_state.feedback = {}
            st.session_state.channel_stats = {
                'Website': 0,
                'WhatsApp': 0,
                'Email': 0,
                'Manual': 0
            }
            
            conn = st.session_state.db_conn
            c = conn.cursor()
            c.execute("DELETE FROM chat_history")
            c.execute("DELETE FROM tickets")
            c.execute("DELETE FROM feedback")
            c.execute("DELETE FROM analytics")
            conn.commit()
            
            st.success("✅ All data cleared!")
            time.sleep(1)
            st.rerun()

# Main page header
st.markdown('<div class="main-header">🤖 AI Customer Support Agent Pro</div>', unsafe_allow_html=True)

# Contact banner
st.markdown(f"""
<div class="contact-banner">
    <h3 style='color: white; text-align: center; margin-bottom: 20px; font-size: 1.8rem;'>
        📞 Get Support Instantly
    </h3>
    <div style='display: flex; justify-content: center; gap: 20px; flex-wrap: wrap;'>
        <div class="contact-card">
            <h4>📧 Email Support</h4>
            <p>{st.session_state.gmail_config['email']}</p>
            <small style='color: #6C757D;'>24/7 Response</small>
        </div>
        <div class="contact-card">
            <h4>💬 WhatsApp Support</h4>
            <p>{st.session_state.whatsapp_config['phone_number']}</p>
            <small style='color: #6C757D;'>Instant Messaging</small>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Channel metrics
st.markdown("### 📊 Channel Overview")
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.metric("🌐 Website", st.session_state.channel_stats.get('Website', 0), delta=None)
with col2:
    st.metric("💬 WhatsApp", st.session_state.channel_stats.get('WhatsApp', 0), delta=None)
with col3:
    st.metric("📧 Email", st.session_state.channel_stats.get('Email', 0), delta=None)
with col4:
    st.metric("👤 Manual", st.session_state.channel_stats.get('Manual', 0), delta=None)

st.markdown("---")

# Tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "💬 Website Chat", 
    "📱 WhatsApp Support", 
    "📧 Email Management",
    "🎫 Ticket System", 
    "📊 Analytics"
])

# Tab 1: Website Chat Agent
with tab1:
    st.markdown('<div class="sub-header">💬 Live Chat Interface</div>', unsafe_allow_html=True)
    st.caption("🌐 Real-time customer support on your website")
    
    if not st.session_state.vector_store:
        st.markdown('<div class="info-box">ℹ️ Please upload and process knowledge base files in the sidebar to enable AI chat.</div>', unsafe_allow_html=True)
    elif not openai_api_key:
        st.markdown('<div class="warning-box">⚠️ Please enter your OpenAI API key in the sidebar to activate the AI agent.</div>', unsafe_allow_html=True)
    else:
        # Chat container
        chat_container = st.container()
        
        with chat_container:
            for i, chat in enumerate(st.session_state.chat_history):
                if chat['role'] == 'user':
                    st.markdown(
                        f'<div class="chat-message user-message">'
                        f'<strong>👤 Customer:</strong><br>{chat["message"]}'
                        f'<br><small style="color: #6C757D;">🌍 {chat.get("language", "English")} | '
                        f'📂 {chat.get("category", "General")}</small>'
                        f'</div>', 
                        unsafe_allow_html=True
                    )
                else:
                    confidence = chat.get('confidence', 0)
                    response_time = chat.get('response_time', 0)
                    confidence_badge = get_confidence_badge(confidence)
                    
                    st.markdown(
                        f'<div class="chat-message bot-message">'
                        f'<strong>🤖 AI Agent:</strong><br>{chat["message"]}<br><br>'
                        f'<div style="display: flex; gap: 10px; flex-wrap: wrap; margin-top: 10px;">'
                        f'{confidence_badge}'
                        f'<span style="color: #6C757D;">⏱️ {response_time:.2f}s</span>'
                        f'</div>'
                        f'</div>', 
                        unsafe_allow_html=True
                    )
                    
                    # Feedback buttons
                    if i not in st.session_state.feedback:
                        col1, col2, col3 = st.columns([1, 1, 10])
                        with col1:
                            if st.button("👍 Helpful", key=f"up_{i}"):
                                st.session_state.feedback[i] = 'positive'
                                st.rerun()
                        with col2:
                            if st.button("👎 Not Helpful", key=f"down_{i}"):
                                st.session_state.feedback[i] = 'negative'
                                st.rerun()
                    else:
                        if st.session_state.feedback[i] == 'positive':
                            st.markdown('<div class="success-box">✓ Marked as helpful - Thank you for your feedback!</div>', unsafe_allow_html=True)
                        else:
                            st.markdown('<div class="warning-box">✗ Marked as not helpful - We\'ll improve our responses</div>', unsafe_allow_html=True)
        
        # Input form
        st.markdown("---")
        with st.form(key="chat_form", clear_on_submit=True):
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input(
                    "Your message:", 
                    placeholder="Type your question here...",
                    label_visibility="collapsed"
                )
            with col2:
                submit_button = st.form_submit_button("Send 🚀", use_container_width=True, type="primary")
        
        if submit_button and user_input:
            start_time = time.time()
            
            # Detect language and category
            language = detect_language(user_input)
            category = categorize_query(user_input)
            
            # Update analytics
            st.session_state.analytics['total_queries'] += 1
            st.session_state.analytics['languages'][language] = st.session_state.analytics['languages'].get(language, 0) + 1
            st.session_state.analytics['categories'][category] += 1
            st.session_state.channel_stats['Website'] += 1
            
            # Save user message
            user_chat = {
                'role': 'user',
                'message': user_input,
                'timestamp': datetime.now(),
                'language': language,
                'category': category
            }
            st.session_state.chat_history.append(user_chat)
            save_chat_to_db(user_chat)
            
            # Get AI response
            with st.spinner(f"🤔 Thinking... (Detected: {language})"):
                answer, confidence, source_docs = get_ai_response(
                    user_input, 
                    st.session_state.vector_store, 
                    openai_api_key,
                    language
                )
            
            response_time = time.time() - start_time
            
            # Check if escalation needed
            if confidence < 0.6:
                ticket_id = create_ticket(user_input, language, category, confidence, 'Website')
                priority = assign_priority(confidence, category)
                escalation_msg = f"\n\n⚠️ **Escalation Notice:** Your query has been forwarded to our support team for personalized assistance.\n\n📋 **Ticket Details:**\n- Ticket ID: `{ticket_id}`\n- Priority: {priority}\n- Expected Response: Within 24 hours"
                if language != 'English':
                    escalation_msg = translate_text(escalation_msg, language)
                answer += escalation_msg
            else:
                st.session_state.analytics['answered'] += 1
            
            # Save bot response
            bot_chat = {
                'role': 'bot',
                'message': answer,
                'timestamp': datetime.now(),
                'confidence': confidence,
                'response_time': response_time,
                'language': language,
                'category': category
            }
            st.session_state.chat_history.append(bot_chat)
            save_chat_to_db(bot_chat)
            
            update_daily_analytics()
            
            st.rerun()

# Tab 2: WhatsApp Support
with tab2:
    st.markdown('<div class="sub-header">📱 WhatsApp Integration</div>', unsafe_allow_html=True)
    st.caption(f"Connected: {st.session_state.whatsapp_config['phone_number']}")
    
    st.markdown("---")
    
    # Quick chat link
    col1, col2 = st.columns([3, 1])
    with col1:
        st.markdown("### 🚀 Quick Chat Link")
        quick_message = st.text_input(
            "Pre-fill message (optional)", 
            placeholder="Hello, I need assistance with..."
        )
    with col2:
        st.write("")
        st.write("")
        wa_link = open_whatsapp_chat(quick_message if quick_message else "Hello! I need support.")
        st.markdown(
            f'<a href="{wa_link}" target="_blank" class="whatsapp-button">💬 Open WhatsApp</a>', 
            unsafe_allow_html=True
        )
    
    st.markdown("---")
    
    # Process incoming messages
    st.markdown("### 📥 Process Customer Queries")
    
    with st.form("whatsapp_form"):
        col1, col2 = st.columns([3, 1])
        with col1:
            wa_message = st.text_area(
                "Customer WhatsApp Message", 
                placeholder="Enter customer query here...", 
                height=120
            )
        with col2:
            wa_phone = st.text_input("Customer Phone", placeholder="+1234567890")
        
        process_wa = st.form_submit_button("🤖 Generate AI Response", type="primary", use_container_width=True)
    
    if process_wa and wa_message and openai_api_key and st.session_state.vector_store:
        with st.spinner("🔄 Processing WhatsApp message..."):
            answer, confidence, ticket_id = process_multi_channel_query(
                wa_message, 
                'WhatsApp', 
                openai_api_key
            )
            
            st.markdown('<div class="success-box">✅ AI Response Generated Successfully!</div>', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**📨 Customer Message**")
                st.info(wa_message)
            with col2:
                st.markdown("**🤖 AI Response**")
                st.success(answer)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Confidence Score", f"{confidence:.1%}")
            with col2:
                st.metric("Status", "Escalated ⚠️" if ticket_id else "Resolved ✅")
            with col3:
                if ticket_id:
                    st.metric("Ticket ID", ticket_id)
            
            # Send response link
            if wa_phone:
                st.markdown("---")
                st.markdown("### 📤 Send Response")
                send_link = create_whatsapp_link(wa_phone, answer)
                st.markdown(
                    f'<a href="{send_link}" target="_blank" class="whatsapp-button">📱 Send to Customer</a>', 
                    unsafe_allow_html=True
                )
            
            # Save interaction
            st.session_state.whatsapp_messages.append({
                'message': wa_message,
                'response': answer,
                'confidence': confidence,
                'ticket_id': ticket_id,
                'phone': wa_phone,
                'timestamp': datetime.now()
            })
    
    # Recent interactions
    if st.session_state.whatsapp_messages:
        st.markdown("---")
        st.markdown("### 📋 Recent WhatsApp Interactions")
        
        for idx, msg in enumerate(reversed(st.session_state.whatsapp_messages[-10:])):
            confidence_badge = get_confidence_badge(msg['confidence'])
            
            with st.expander(
                f"💬 {msg['timestamp'].strftime('%Y-%m-%d %H:%M')} | Phone: {msg.get('phone', 'N/A')[:15]}"
            ):
                st.markdown(f"**Confidence:** {confidence_badge}", unsafe_allow_html=True)
                st.markdown(f"**Customer:** {msg['message']}")
                st.markdown(f"**AI Response:** {msg['response']}")
                
                if msg['ticket_id']:
                    st.markdown(f'<div class="warning-box">⚠️ Escalated - Ticket: {msg["ticket_id"]}</div>', unsafe_allow_html=True)
                
                if msg.get('phone'):
                    resend_link = create_whatsapp_link(msg['phone'], msg['response'])
                    st.markdown(
                        f'<a href="{resend_link}" target="_blank" class="whatsapp-button">🔄 Resend Response</a>', 
                        unsafe_allow_html=True
                    )

# Tab 3: Email Management
with tab3:
    st.markdown('<div class="sub-header">📧 Email Auto-Response System</div>', unsafe_allow_html=True)
    st.caption(f"Connected: {st.session_state.gmail_config['email']}")
    
    # Gmail configuration
    with st.expander("⚙️ Gmail Setup Guide", expanded=False):
        st.markdown("""
        <div class="info-box">
        <strong>📝 Setup Instructions:</strong><br>
        1. Go to your Google Account settings<br>
        2. Enable 2-Step Verification<br>
        3. Navigate to Security → App Passwords<br>
        4. Generate a new App Password for 'Mail'<br>
        5. Enter the 16-character password below
        </div>
        """, unsafe_allow_html=True)
        
        gmail_app_password = st.text_input("Gmail App Password", type="password", key="gmail_pwd")
        
        if gmail_app_password:
            st.markdown('<div class="success-box">✅ Gmail credentials configured successfully!</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Check inbox
    st.markdown("### 📬 Check Inbox")
    
    col1, col2 = st.columns([3, 1])
    with col1:
        max_emails = st.slider("Number of emails to check", 1, 20, 10)
    with col2:
        st.write("")
        st.write("")
        check_inbox = st.button("🔄 Refresh Inbox", type="primary")
    
    if check_inbox and gmail_app_password:
        with st.spinner("🔍 Checking Gmail inbox..."):
            new_emails = check_gmail_inbox(gmail_app_password, max_emails)
            
            if new_emails:
                st.markdown(f'<div class="success-box">✅ Found {len(new_emails)} unread email(s)</div>', unsafe_allow_html=True)
                
                for email_data in new_emails:
                    with st.expander(f"📧 {email_data['from'][:50]} - {email_data['subject'][:50]}"):
                        st.markdown(f"**Subject:** {email_data['subject']}")
                        st.markdown(f"**From:** {email_data['from']}")
                        st.markdown(f"**Preview:** {email_data['body'][:300]}...")
                        
                        if st.button(f"🤖 Generate Response", key=f"respond_{email_data['id']}"):
                            with st.spinner("⚙️ Generating response..."):
                                answer, confidence, ticket_id = process_multi_channel_query(
                                    email_data['body'], 
                                    'Email', 
                                    openai_api_key
                                )
                                
                                st.markdown('<div class="success-box">✅ Response generated!</div>', unsafe_allow_html=True)
                                st.write(answer)
                                
                                st.session_state.email_queue.append({
                                    'from': email_data['from'],
                                    'subject': email_data['subject'],
                                    'body': email_data['body'],
                                    'response': answer,
                                    'confidence': confidence,
                                    'ticket_id': ticket_id,
                                    'timestamp': datetime.now(),
                                    'sent': False
                                })
            else:
                st.markdown('<div class="info-box">ℹ️ No new unread emails found</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Manual email composition
    st.markdown("### ✍️ Compose Email Response")
    
    with st.form("email_form"):
        customer_email = st.text_input("Customer Email", placeholder="customer@example.com")
        email_subject = st.text_input("Subject", placeholder="Re: Your inquiry")
        email_body = st.text_area("Customer's Email", placeholder="Customer query...", height=150)
        
        process_email = st.form_submit_button("🤖 Generate AI Response", type="primary", use_container_width=True)
    
    if process_email and email_body and openai_api_key and st.session_state.vector_store:
        with st.spinner("⚙️ Processing email..."):
            answer, confidence, ticket_id = process_multi_channel_query(
                email_body, 
                'Email', 
                openai_api_key
            )
            
            email_response = f"""Dear Valued Customer,

Thank you for reaching out to us. Here is the information regarding your inquiry:

{answer}

Best regards,
Customer Support Team
{st.session_state.gmail_config['email']}

---
This is an automated AI response. For further assistance, please reply to this email or contact our support team directly.
"""
            
            if ticket_id:
                email_response += f"\n\nReference Ticket ID: {ticket_id}"
            
            st.markdown('<div class="success-box">✅ Email Response Generated Successfully!</div>', unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**📨 Customer Email**")
                st.info(email_body)
            with col2:
                st.markdown("**🤖 AI Response**")
                st.success(answer)
            
            st.markdown("**📧 Complete Email Draft**")
            st.code(email_response, language="text")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Confidence", f"{confidence:.1%}")
            with col2:
                st.metric("Status", "Escalated ⚠️" if ticket_id else "Resolved ✅")
            with col3:
                if ticket_id:
                    st.metric("Ticket", ticket_id)
            
            # Send email
            if gmail_app_password and customer_email:
                st.markdown("---")
                if st.button("📤 Send Email Now", type="primary", use_container_width=True):
                    response_subject = f"Re: {email_subject}" if email_subject else "Response to your inquiry"
                    
                    with st.spinner("📧 Sending email..."):
                        if send_gmail(customer_email, response_subject, email_response, gmail_app_password):
                            st.markdown('<div class="success-box">✅ Email sent successfully!</div>', unsafe_allow_html=True)
                            st.balloons()
            
            st.session_state.email_queue.append({
                'from': customer_email,
                'subject': email_subject,
                'body': email_body,
                'response': answer,
                'confidence': confidence,
                'ticket_id': ticket_id,
                'timestamp': datetime.now(),
                'sent': False
            })
    
    # Email queue
    if st.session_state.email_queue:
        st.markdown("---")
        st.markdown("### 📬 Email Queue")
        
        email_df = pd.DataFrame([
            {
                'Time': email['timestamp'].strftime('%Y-%m-%d %H:%M'),
                'From': email['from'][:30],
                'Subject': email['subject'][:40],
                'Confidence': f"{email['confidence']:.1%}",
                'Status': '⚠️ Escalated' if email['ticket_id'] else '✅ Resolved',
                'Sent': '✅ Yes' if email.get('sent') else '⏳ Pending'
            }
            for email in reversed(st.session_state.email_queue[-20:])
        ])
        
        st.dataframe(email_df, use_container_width=True, hide_index=True)

# Tab 4: Ticket Management
with tab4:
    st.markdown('<div class="sub-header">🎫 Support Ticket System</div>', unsafe_allow_html=True)
    
    if st.session_state.tickets:
        # Filters
        col1, col2, col3 = st.columns(3)
        with col1:
            status_filter = st.multiselect(
                "Filter by Status",
                options=['Open', 'In Progress', 'Closed'],
                default=['Open', 'In Progress']
            )
        with col2:
            priority_filter = st.multiselect(
                "Filter by Priority",
                options=['High', 'Medium', 'Low'],
                default=['High', 'Medium', 'Low']
            )
        with col3:
            category_filter = st.multiselect(
                "Filter by Category",
                options=['Billing', 'Technical', 'General', 'Product'],
                default=['Billing', 'Technical', 'General', 'Product']
            )
        
        df_tickets = pd.DataFrame(st.session_state.tickets)
        
        filtered_df = df_tickets[
            (df_tickets['status'].isin(status_filter)) &
            (df_tickets['priority'].isin(priority_filter)) &
            (df_tickets['category'].isin(category_filter))
        ]
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Tickets", len(df_tickets))
        with col2:
            st.metric("Open", len(df_tickets[df_tickets['status'] == 'Open']))
        with col3:
            st.metric("High Priority", len(df_tickets[df_tickets['priority'] == 'High']))
        with col4:
            avg_resolution = df_tickets[df_tickets['resolution_time'].notna()]['resolution_time'].mean()
            st.metric("Avg Resolution", f"{avg_resolution:.1f}h" if not pd.isna(avg_resolution) else "N/A")
        
        st.markdown("---")
        
        # Ticket cards
        for idx, ticket in filtered_df.iterrows():
            priority_badge = get_priority_badge(ticket['priority'])
            status_badge = get_status_badge(ticket['status'])
            
            with st.expander(
                f"Ticket #{ticket['id']} | {ticket['category']} | {ticket['channel']}", 
                expanded=False
            ):
                st.markdown(f"""
                <div style='margin-bottom: 1rem;'>
                    {priority_badge} {status_badge}
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown(f"**Query:** {ticket['query']}")
                    st.markdown(f"**Category:** {ticket['category']}")
                    st.markdown(f"**Language:** {ticket['language']}")
                    st.markdown(f"**Channel:** {ticket['channel']}")
                with col2:
                    st.markdown(f"**Assigned To:** {ticket['assigned_to']}")
                    st.markdown(f"**Created:** {ticket['timestamp']}")
                    if ticket['resolved_at']:
                        st.markdown(f"**Resolved:** {ticket['resolved_at']}")
                        st.markdown(f"**Resolution Time:** {ticket['resolution_time']:.1f}h")
                
                st.markdown("---")
                
                col1, col2, col3 = st.columns([2, 2, 1])
                with col1:
                    new_status = st.selectbox(
                        "Update Status",
                        options=['Open', 'In Progress', 'Closed'],
                        index=['Open', 'In Progress', 'Closed'].index(ticket['status']),
                        key=f"status_{ticket['id']}"
                    )
                with col2:
                    new_assignee = st.selectbox(
                        "Reassign To",
                        options=['Agent A', 'Agent B', 'Agent C', 'Agent D'],
                        index=['Agent A', 'Agent B', 'Agent C', 'Agent D'].index(ticket['assigned_to']),
                        key=f"assign_{ticket['id']}"
                    )
                with col3:
                    st.write("")
                    st.write("")
                    if st.button("💾 Update", key=f"update_{ticket['id']}", type="primary"):
                        # Calculate resolution time if closing
                        if new_status == 'Closed' and ticket['status'] != 'Closed':
                            created = datetime.strptime(ticket['timestamp'], "%Y-%m-%d %H:%M:%S")
                            resolution_time = (datetime.now() - created).total_seconds() / 3600
                            st.session_state.tickets[idx]['resolved_at'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                            st.session_state.tickets[idx]['resolution_time'] = resolution_time
                        
                        st.session_state.tickets[idx]['status'] = new_status
                        st.session_state.tickets[idx]['assigned_to'] = new_assignee
                        
                        st.markdown('<div class="success-box">✅ Ticket updated successfully!</div>', unsafe_allow_html=True)
                        time.sleep(1)
                        st.rerun()
    else:
        st.markdown('<div class="info-box">ℹ️ No tickets have been created yet. Tickets are automatically generated when AI confidence is low.</div>', unsafe_allow_html=True)

# Tab 5: Analytics Dashboard
with tab5:
    st.markdown('<div class="sub-header">📊 Performance Analytics</div>', unsafe_allow_html=True)
    
    analytics = st.session_state.analytics
    
    # Key metrics
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Queries", analytics.get('total_queries', 0))
    with col2:
        st.metric("Answered", analytics.get('answered', 0), delta="Auto-resolved")
    with col3:
        st.metric("Escalated", analytics.get('escalated', 0), delta="Human review")
    with col4:
        total = analytics.get('total_queries', 0)
        answered = analytics.get('answered', 0)
        resolution_rate = (answered / total * 100) if total > 0 else 0
        st.metric("Resolution Rate", f"{resolution_rate:.1f}%")
    
    st.markdown("---")
    
    # Query trends
    if st.session_state.chat_history:
        st.markdown("### 📈 Query Trends Over Time")
        
        chat_df = pd.DataFrame([
            {
                'timestamp': chat['timestamp'],
                'role': chat['role'],
                'confidence': chat.get('confidence', None),
                'response_time': chat.get('response_time', None)
            }
            for chat in st.session_state.chat_history
        ])
        
        chat_df['hour'] = pd.to_datetime(chat_df['timestamp']).dt.floor('H')
        hourly_queries = chat_df[chat_df['role'] == 'user'].groupby('hour').size().reset_index(name='queries')
        
        if len(hourly_queries) > 0:
            fig_timeline = px.line(
                hourly_queries,
                x='hour',
                y='queries',
                title="Customer Queries Timeline",
                labels={'hour': 'Time', 'queries': 'Number of Queries'},
                markers=True
            )
            fig_timeline.update_traces(line_color='#0066CC', line_width=3)
            fig_timeline.update_layout(
                plot_bgcolor='white',
                paper_bgcolor='white',
                font=dict(color='#212529')
            )
            st.plotly_chart(fig_timeline, use_container_width=True)
    
    st.markdown("---")
    
    # Charts
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📊 Query Distribution")
        answered = analytics.get('answered', 0)
        escalated = analytics.get('escalated', 0)
        
        fig_queries = go.Figure(data=[
            go.Bar(
                x=['Auto-Resolved', 'Escalated'],
                y=[answered, escalated],
                marker_color=['#00A86B', '#DC143C'],
                text=[answered, escalated],
                textposition='auto',
                textfont=dict(size=14, color='white')
            )
        ])
        fig_queries.update_layout(
            title="Resolution Status",
            xaxis_title="Status",
            yaxis_title="Count",
            height=350,
            plot_bgcolor='white',
            paper_bgcolor='white',
            font=dict(color='#212529')
        )
        st.plotly_chart(fig_queries, use_container_width=True)
    
    with col2:
        st.markdown("### 📱 Channel Distribution")
        channel_data = {k: v for k, v in st.session_state.channel_stats.items() if v > 0}
        if channel_data:
            colors = ['#0066CC', '#25D366', '#EA4335', '#6C757D']
            fig_channel = px.pie(
                values=list(channel_data.values()),
                names=list(channel_data.keys()),
                title="Queries by Channel",
                color_discrete_sequence=colors
            )
            fig_channel.update_layout(
                height=350,
                paper_bgcolor='white',
                font=dict(color='#212529')
            )
            st.plotly_chart(fig_channel, use_container_width=True)
        else:
            st.markdown('<div class="info-box">ℹ️ No channel data available yet</div>', unsafe_allow_html=True)
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📂 Category Breakdown")
        categories = analytics.get('categories', {})
        category_df = pd.DataFrame({
            'Category': list(categories.keys()),
            'Count': list(categories.values())
        })
        
        fig_category = px.bar(
            category_df,
            x='Category',
            y='Count',
            color='Category',
            title="Queries by Category",
            color_discrete_sequence=['#0066CC', '#00A86B', '#FF8C00', '#DC143C']
        )
        fig_category.update_layout(
            height=350,
            plot_bgcolor='white',
            paper_bgcolor='white',
            font=dict(color='#212529'),
            showlegend=False
        )
        st.plotly_chart(fig_category, use_container_width=True)
    
    with col2:
        st.markdown("### ⚡ Performance Metrics")
        bot_messages = [msg for msg in st.session_state.chat_history if msg['role'] == 'bot']
        
        if bot_messages:
            avg_response_time = sum(msg.get('response_time', 0) for msg in bot_messages) / len(bot_messages)
            avg_confidence = sum(msg.get('confidence', 0) for msg in bot_messages) / len(bot_messages)
            
            fig_perf = go.Figure(data=[
                go.Bar(
                    x=['Avg Response Time (s)', 'Avg Confidence'],
                    y=[avg_response_time, avg_confidence],
                    marker_color=['#0066CC', '#00A86B'],
                    text=[f"{avg_response_time:.2f}s", f"{avg_confidence:.1%}"],
                    textposition='auto',
                    textfont=dict(size=14, color='white')
                )
            ])
            fig_perf.update_layout(
                title="System Performance",
                height=350,
                yaxis_title="Value",
                plot_bgcolor='white',
                paper_bgcolor='white',
                font=dict(color='#212529')
            )
            st.plotly_chart(fig_perf, use_container_width=True)
        else:
            st.markdown('<div class="info-box">ℹ️ No performance data available yet</div>', unsafe_allow_html=True)

# Footer
st.markdown("---")
st.markdown(f"""
<div class="footer">
    <h3 style='color: var(--primary-dark); margin-bottom: 1rem;'>🤖 AI Customer Support Agent Pro v4.0</h3>
    <p style='font-weight: 600; margin-bottom: 0.5rem;'>Powered by OpenAI GPT-3.5, LangChain & FAISS</p>
    <p style='margin-bottom: 1rem;'>
        📧 {st.session_state.gmail_config['email']} | 
        💬 {st.session_state.whatsapp_config['phone_number']}
    </p>
    <div style='background: var(--bg-light); padding: 1rem; border-radius: 8px; margin-top: 1rem;'>
        <p style='margin: 0; font-size: 0.9rem; color: var(--text-dark);'>
            <strong>✅ Features:</strong> Multi-Channel Support (Website • WhatsApp • Email) | 
            AI-Powered Responses | Document Processing (PDF • DOCX • Excel • Images) | 
            OCR Technology | Web Scraping | Multilingual Support | Smart Ticket Management | 
            Real-time Analytics | Semantic Search | Auto-Escalation | Performance Tracking
        </p>
    </div>
</div>
""", unsafe_allow_html=True)
